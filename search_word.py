# الملف: search_word.py
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from PIL import Image, ImageTk
import os, sys

# ---------------- تحديد المسار الأساسي للصور والأيقونة ----------------  
if getattr(sys, 'frozen', False):
    base_path = sys._MEIPASS
else:
    base_path = os.path.dirname(__file__)

def get_image_path(filename):
    return os.path.join(base_path, "images", filename)

def open_word_search_window(root):
    root.withdraw()  # إخفاء الفورم الرئيسي

    # --- نافذة البحث ---
    search_window = tk.Toplevel(root)
    search_window.title("🔍 استعراض المهام")
    search_window.state('zoomed')

    # ألوان متوافقة مع الفورم الرئيسي
    form_bg = '#D35400'
    button_colors = ['#1E90FF', '#28a745', '#FF4500', '#8A2BE2', '#FFD700', '#00CED1']
    button_fg = "white"
    search_window.configure(bg=form_bg)

    # ---------------- أيقونة النافذة ----------------
    icon_path = get_image_path("ITP1.ico")
    if os.path.exists(icon_path):
        search_window.iconbitmap(icon_path)

    # ---------------- منع الإغلاق بزر X ----------------
    def on_close():
        messagebox.showinfo("تنبيه", "استخدم زر العودة للقائمة الأساسية لإغلاق النافذة بشكل صحيح.")
    search_window.protocol("WM_DELETE_WINDOW", on_close)

    selected_file = tk.StringVar()

    def open_file():
        file_path = filedialog.askopenfilename(
            title="اختر ملف Word",
            filetypes=[("Word Files", "*.docx")]
        )
        if file_path:
            selected_file.set(file_path)

    # --------------------- إعداد صورة loup المتحركة ---------------------
    loup_label = None
    loup_img_path = get_image_path("loup.jpg")  # ضع صورة العدسة هنا
    if os.path.exists(loup_img_path):
        loup_img = ImageTk.PhotoImage(Image.open(loup_img_path).resize((30, 30)))
        loup_label = tk.Label(search_window, image=loup_img, bg=form_bg)
        loup_label.place_forget()  # أخفيها مبدئيًا

    moving = False  # حالة حركة الصورة

    def move_loup():
        if not moving or loup_label is None:
            return
        x = loup_label.winfo_x() + 5
        if x > search_window.winfo_width() - 40:  # إعادة للصفر عند الوصول لليمين
            x = 10
        loup_label.place(x=x, y=search_btn.winfo_y() - 35)
        search_window.after(50, move_loup)

    # --------------------- دالة البحث ---------------------
    def search_in_word():
        nonlocal moving
        if loup_label:
            moving = True
            loup_label.place(x=10, y=search_btn.winfo_y() - 35)
            move_loup()

        search_window.update_idletasks()  # تحديث الفورم قبل بدء القراءة

        file_path = selected_file.get().strip()
        inspector = inspector_entry.get().strip()
        month = month_combobox.get().strip()
        listbox.delete(0, tk.END)

        if not file_path:
            messagebox.showwarning("تنبيه", "⚠️ الرجاء اختيار الملف أولاً.")
            if loup_label: moving = False; loup_label.place_forget()
            return
        if not inspector:
            messagebox.showwarning("تنبيه", "⚠️ الرجاء إدخال اسم المفتش.")
            if loup_label: moving = False; loup_label.place_forget()
            return
        if not month:
            messagebox.showwarning("تنبيه", "⚠️ الرجاء اختيار الشهر.")
            if loup_label: moving = False; loup_label.place_forget()
            return

        try:
            doc = Document(file_path)
            inspector_found = any(inspector in para.text for para in doc.paragraphs)
            if not inspector_found:
                messagebox.showinfo("❌ لا توجد نتائج", f"لم يتم العثور على المفتش '{inspector}' في الملف.")
                if loup_label: moving = False; loup_label.place_forget()
                return

            found_tasks = set()
            for table_idx, table in enumerate(doc.tables, start=1):
                last_month_value = ""
                for row in table.rows:
                    month_cell = row.cells[2].text.strip()
                    if month_cell:
                        last_month_value = month_cell
                    else:
                        month_cell = last_month_value

                    task_cell = row.cells[4].text.strip()
                    if month in month_cell and task_cell:
                        current_line = ""
                        for char in task_cell:
                            if char.isalnum() or char.isspace():
                                current_line += char
                            else:
                                if current_line.strip() and current_line.strip() not in found_tasks:
                                    listbox.insert(tk.END, current_line.strip())
                                    found_tasks.add(current_line.strip())
                                current_line = ""
                        if current_line.strip() and current_line.strip() not in found_tasks:
                            listbox.insert(tk.END, current_line.strip())
                            found_tasks.add(current_line.strip())

            if found_tasks:
                messagebox.showinfo("✅ النتيجة", f"تم العثور على نتائج للمفتش '{inspector}' والشهر '{month}'.")
            else:
                messagebox.showinfo("❌ لا توجد نتائج", f"تم العثور على المفتش '{inspector}' ولكن لا توجد مهام تطابق الشهر '{month}'.")

        except Exception as e:
            messagebox.showerror("❌ خطأ", f"حدث خطأ أثناء القراءة: {e}")

        if loup_label:
            moving = False
            loup_label.place_forget()  # إخفاء الصورة عند انتهاء البحث

    # ---------------- واجهة العناصر ----------------
    button01 = tk.Button(search_window, text="📂 اختيار ملف Word", command=open_file,
                         bg=button_colors[4], fg=button_fg, font=("Arial", 12, "bold"))
    button01.pack(pady=10)

    file_label = tk.Label(search_window, textvariable=selected_file, bg=form_bg,
                          fg=button_fg, font=("Arial", 10, "italic"), wraplength=850, justify="right")
    file_label.pack()

    inspector_label = tk.Label(search_window, text="🧑 أدخل اسم المفتش:", bg=form_bg, fg=button_fg)
    inspector_label.pack(pady=5)
    inspector_entry = tk.Entry(search_window, font=("Arial", 12, "bold"), justify='right')
    inspector_entry.pack()

    month_label = tk.Label(search_window, text="🗓️ اختر الشهر:", bg=form_bg, fg=button_fg)
    month_label.pack(pady=5)
    month_combobox = ttk.Combobox(
        search_window,
        values=["جانفي", "فيفري", "مارس", "أفريل", "ماي", "جوان",
                "جويلية", "أوت", "سبتمبر", "أكتوبر", "نوفمبر", "ديسمبر"],
        justify='right'
    )
    month_combobox.pack()

    search_btn = tk.Button(search_window, text="🔍 بحث", command=search_in_word,
                           bg=button_colors[1], fg=button_fg, font=("Arial", 12, "bold"))
    search_btn.pack(pady=10)

    # ---------------- Listbox لعرض النتائج ----------------
    listbox = tk.Listbox(
        search_window,
        width=120,
        height=20,
        font=("Arial", 12, "bold"),
        bg="#FFFFFF",
        fg="#000000",
        justify='right'
    )
    listbox.pack(pady=10, padx=10, fill='both', expand=True)

    # ---------------- زر العودة للقائمة الرئيسية ----------------
    back_btn = tk.Button(search_window, text="⬅️ العودة للقائمة الأساسية",
                         command=lambda: [search_window.destroy(), root.deiconify()],
                         bg=button_colors[0], fg=button_fg, font=("Arial", 12, "bold"))
    back_btn.pack(pady=10)

    # الاحتفاظ بالصور
    if loup_label:
        search_window.loup_img = loup_img
