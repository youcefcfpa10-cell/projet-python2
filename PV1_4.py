import tkinter as tk
from tkinter import messagebox
import os
import json
from datetime import datetime
from docx import Document
from docx2pdf import convert
from notes_window import open_notes_window
import PFP  # لإرجاع المستخدم إلى القائمة الرئيسية

# --- دالة لحفظ البيانات في JSON ---
def save_to_json(data, filename="data.json"):
    with open(filename, "w", encoding="utf-8-sig") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

def open_pv1_4():
    root = tk.Tk()
    root.title("أعضاء اللجنة - النتائج النهائية")
    root.geometry("1050x450")
    root.configure(bg="#D35400")
    root.option_add('*Font', 'Arial 12')
    root.option_add('*justify', 'right')

    # تعطيل زر X وإظهار تنبيه
    def on_close():
        messagebox.showinfo("تنبيه ⚠️", "لا يمكن إغلاق النافذة مباشرة! الرجاء استخدام زر العودة إلى القائمة الرئيسية.")
    root.protocol("WM_DELETE_WINDOW", on_close)

    # إعداد أيقونة البرنامج
    icon_path = os.path.join("images", "ITP1.ICO")
    if os.path.exists(icon_path):
        root.iconbitmap(icon_path)

    # --- تحميل البيانات من data.json ---
    total, total1 = 0, 0
    all_data = {}
    if os.path.exists("data.json"):
        with open("data.json", "r", encoding="utf-8-sig") as f:
            try:
                all_data = json.load(f)
                total = float(all_data.get("PV1_2", {}).get("total", 0))
                total1 = float(all_data.get("PV1_3", {}).get("total1", 0))
            except Exception as e:
                print("خطأ في تحميل JSON:", e)

    moyenne = round((total + total1) / 3, 2)
    decision = "ناجح" if moyenne >= 10 else "راسب"

    # --- الإطار العلوي للنتائج ---
    top_frame = tk.Frame(root, bg="#D35400")
    top_frame.pack(pady=10)

    labels_text = ["إلقاء الدرس ", " المحادثة", "المجموع", "المعدل", "القرار", "التاريخ"]
    td = round(total + total1, 2)
    values = [total, total1, td, moyenne, decision, datetime.now().strftime("%Y-%m-%d")]
    text_boxes = []

    for i in range(6):
        lbl = tk.Label(top_frame, text=labels_text[i],
                       bg="#A64B00", fg="white", font=("Arial", 13, "bold"),
                       width=15, pady=6, relief="ridge")
        lbl.grid(row=0, column=i, padx=6, pady=4, sticky="nsew")

        txt = tk.Entry(top_frame, justify="center", font=("Arial", 12, "bold"), width=15)
        txt.grid(row=1, column=i, padx=6, pady=4)

        if isinstance(values[i], float):
            txt.insert(0, f"{values[i]:.2f}")
        else:
            txt.insert(0, values[i])

        text_boxes.append(txt)

    date_entry = text_boxes[5]

    # --- إطار الجدول ---
    table_frame = tk.Frame(root, bg="#D35400")
    table_frame.pack(anchor="e", padx=30, pady=(30, 8))

    headers = [("الوظيفة", 0), ("الاسم", 1), ("اللقب", 2), ("العضو", 3)]
    for text, col in headers:
        lbl = tk.Label(table_frame, text=text, bg="#A64B00", fg="white",
                       font=("Arial", 13, "bold"), width=20, pady=6, relief="ridge")
        lbl.grid(row=0, column=col, padx=2, pady=2, sticky="nsew")

    members = ["مفتش المقاطعة", "مدير المؤسسة", "العضو الأول", "العضو الثاني"]
    entries_lname, entries_fname, entries_job = [], [], []

    for i in range(4):
        job_e = tk.Entry(table_frame, justify='right', width=25)
        job_e.grid(row=i + 1, column=0, padx=2, pady=4, sticky="nsew")

        fname_e = tk.Entry(table_frame, justify='right', width=25)
        fname_e.grid(row=i + 1, column=1, padx=2, pady=4, sticky="nsew")

        lname_e = tk.Entry(table_frame, justify='right', width=25)
        lname_e.grid(row=i + 1, column=2, padx=2, pady=4, sticky="nsew")

        member_lbl = tk.Label(table_frame, text=members[i], bg="#FFF8F0", width=12, relief="ridge")
        member_lbl.grid(row=i + 1, column=3, padx=2, pady=4, sticky="nsew")

        entries_job.append(job_e)
        entries_fname.append(fname_e)
        entries_lname.append(lname_e)

    # --- الأزرار ---
    buttons_frame = tk.Frame(root, bg="#D35400")
    buttons_frame.pack(fill='x', padx=20, pady=(40, 12))

    def save_data():
        td = round(total + total1, 2)

        data = {
            "PV1_4": {
                "total": round(total, 2),
                "total1": round(total1, 2),
                "td": td,
                "moyenne": round(moyenne, 2),
                "decision": decision,
                "date": date_entry.get(),
                "members": []
            }
        }

        for i in range(4):
            lname_val = entries_lname[i].get()
            fname_val = entries_fname[i].get()
            job_val = entries_job[i].get()
            data["PV1_4"]["members"].append({
                "member": members[i],
                "lname": lname_val,
                "fname": fname_val,
                "job": job_val
            })

        all_data["PV1_4"] = data["PV1_4"]
        save_to_json(all_data)
        messagebox.showinfo("تم الحفظ ✅", "تم حفظ جميع البيانات بنجاح.")

    def return_to_main_menu():
        root.destroy()
        PFP.open_confirmation_window()

    def fill_all_templates():
        template_path = "templates/PFP.docx"
        if not os.path.exists(template_path):
            messagebox.showerror("خطأ", f"لم يتم العثور على المحضر:\n{template_path}")
            return

        doc = Document(template_path)
        result = messagebox.askquestion("تأكيد الحفظ", "هل أنت متأكد من حفظ جميع البيانات؟", icon='warning')
        if result == 'no':
            return

        replacements = {}

        # ✅ إضافة X أو ☐ حسب القرار
        decision_local = "ناجح" if moyenne >= 10 else "راسب"
        if decision_local == "ناجح":
            replacements["{{pass_exam}}"] = "☒"
            replacements["{{fail_exam}}"] = "☐"
        else:
            replacements["{{pass_exam}}"] = "☐"
            replacements["{{fail_exam}}"] = "☒"

        # ✅ تمرير قيمة td
        if "PV1_4" in all_data:
            replacements["{{td}}"] = str(all_data["PV1_4"].get("td", ""))

        # --- بيانات JSON ---
        for section, section_data in all_data.items():
            if isinstance(section_data, dict):
                for key, val in section_data.items():
                    if isinstance(val, list):
                        for i, v in enumerate(val, start=1):
                            replacements[f"{{{{{key}{i}}}}}"] = str(v)
                            replacements[f"{{{{{key}_{i}}}}}"] = str(v)
                    else:
                        replacements[f"{{{{{key}}}}}"] = str(val)

        # --- أعضاء اللجنة ---
        if "PV1_4" in all_data and "members" in all_data["PV1_4"]:
            for i, m in enumerate(all_data["PV1_4"]["members"], start=1):
                replacements[f"{{{{fname{i}}}}}"] = m.get("fname", "")
                replacements[f"{{{{lname{i}}}}}"] = m.get("lname", "")
                replacements[f"{{{{job{i}}}}}"] = m.get("job", "")

        def replace_text_in_paragraph(paragraph):
            full_text = "".join(run.text for run in paragraph.runs)
            new_text = full_text
            for k, v in replacements.items():
                if k in new_text:
                    new_text = new_text.replace(k, v)
            if new_text != full_text:
                for _ in range(len(paragraph.runs) - 1):
                    paragraph.runs[1].clear()
                    paragraph._element.remove(paragraph.runs[1]._element)
                paragraph.runs[0].text = new_text

        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph)

        teacher_last = all_data.get("Teacher", {}).get("last_name", "").strip()
        teacher_first = all_data.get("Teacher", {}).get("first_name", "").strip()
        output_filename = f"{teacher_last}_{teacher_first}.docx"
        documents_path = os.path.join(os.path.expanduser("~"), "Documents")
        final_folder = os.path.join(documents_path, "محاضر التثبيت", "أساتذة التكوين المهنيPFP")
        os.makedirs(final_folder, exist_ok=True)
        output_path = os.path.join(final_folder, output_filename)

        if os.path.exists(output_path):
            os.remove(output_path)
        doc.save(output_path)

        try:
            pdf_path = output_path.replace(".docx", ".pdf")
            convert(output_path, pdf_path)
        except Exception as e:
            messagebox.showwarning("⚠️ لم يتم إنشاء PDF", f"تم حفظ Word فقط.\nالسبب:\n{e}")

        messagebox.showinfo("✅ تم الحفظ بنجاح",
                            f"تم حفظ الملف في المجلد:\n{final_folder}")

    tk.Button(buttons_frame, text="💾 حفظ البيانات", bg="#4CAF50", fg="white",
              font=("Arial", 12, "bold"), width=15, command=save_data).pack(side="left", padx=10, expand=True)
    tk.Button(buttons_frame, text="🧾 تعبئة البيانات في المحضر    ", bg="#1E90FF", fg="white",
              font=("Arial", 12, "bold"), width=22,
              command=lambda: open_notes_window(root, all_data, fill_all_templates, icon_path)).pack(side="left", padx=10, expand=True)
    tk.Button(buttons_frame, text="⬅️ العودة إلى القائمة الرئيسية", bg="#FF5733", fg="white",
              font=("Arial", 12, "bold"), width=22, command=return_to_main_menu).pack(side="left", padx=10, expand=True)

    root.mainloop()

if __name__ == "__main__":
    open_pv1_4()
