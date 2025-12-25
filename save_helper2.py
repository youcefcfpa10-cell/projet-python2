# save_helper1.py
import json
import os
from docx import Document

# ---------------- حفظ البيانات في ملف JSON ----------------
def save_to_json(form_name, data, file_path="data.json"):
    """
    يحفظ البيانات الخاصة بفورم معين في ملف JSON واحد.
    form_name: اسم الفورم (مثل PV1_4)
    data: القيم المراد حفظها كـ dict
    file_path: مسار ملف JSON
    """
    all_data = {}
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            try:
                all_data = json.load(f)
            except Exception:
                all_data = {}

    all_data[form_name] = data

    with open(file_path, "w", encoding="utf-8") as f:
        json.dump(all_data, f, ensure_ascii=False, indent=4)


# ---------------- استرجاع البيانات من JSON ----------------
def load_from_json(file_path="data.json"):
    """
    يعيد كل البيانات المخزنة في ملف JSON.
    """
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            try:
                return json.load(f)
            except Exception:
                return {}
    return {}


# ---------------- تعبئة القالب Word ----------------
def fill_template_with_check(template_path, output_path, replacements):
    """
    template_path: مسار قالب Word الأصلي
    output_path: مسار حفظ القالب بعد تعبئته
    replacements: dict يحتوي على {علامة: قيمة} لتبديل النصوص
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"لم يتم العثور على القالب: {template_path}")

    doc = Document(template_path)

    def replace_text_in_paragraph(paragraph):
        for key, val in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    run.text = run.text.replace(key, val)

    # استبدال النصوص في كل الفقرات
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph)

    # استبدال النصوص داخل الجداول
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
