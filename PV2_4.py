import tkinter as tk
from tkinter import messagebox
import tkinter.font as tkFont
import os
import json
from datetime import datetime
from docx import Document
from docx2pdf import convert
import ADJOIN

# ---------------- حفظ JSON ----------------
def save_to_json(data, filename="data1.json"):
    with open(filename, "w", encoding="utf-8-sig") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

# ================== 🔴 التعديلات الخاصة بتسطيح JSON ==================

# -------- تفكيك JSON (يدعم dict + list) --------
def flatten_json(data, parent_key=""):
    items = {}
    if isinstance(data, dict):
        for k, v in data.items():
            new_key = f"{parent_key}.{k}" if parent_key else k
            items.update(flatten_json(v, new_key))
    elif isinstance(data, list):
        for i, v in enumerate(data):
            new_key = f"{parent_key}.{i}"
            items.update(flatten_json(v, new_key))
    else:
        items[parent_key] = str(data)
    return items

# -------- دمج جميع الأقسام في الجذر للقالب --------
def flatten_all_sections(data):
    flat_data = {}
    for section in data.values():
        if isinstance(section, dict):
            flat_data.update(section)
        elif isinstance(section, list):
            # إذا كان القسم قائمة، نفرغها بعلامة مفلطحة
            for i, item in enumerate(section):
                if isinstance(item, dict):
                    for k, v in item.items():
                        flat_data[f"{k}_{i}"] = str(v)
                else:
                    flat_data[f"item_{i}"] = str(item)
    return flat_data

# -------- تعويض ذكي داخل Word (runs + tables) --------
def replace_in_doc(doc, data):
    flat_data = flatten_json(data)

    def replace_paragraphs(paragraphs):
        for paragraph in paragraphs:
            for run in paragraph.runs:
                for key, value in flat_data.items():
                    tag = "{{" + key + "}}"
                    if tag in run.text:
                        run.text = run.text.replace(tag, value)

    replace_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_paragraphs(cell.paragraphs)

# ================== 🔴 التعديل ينتهي هنا ==================

# ---------------- النافذة الرئيسية ----------------
def open_pv2_4():
    root = tk.Tk()
    root.title("أعضاء اللجنة - النتائج النهائية")
    root.geometry("1150x500")
    root.configure(bg="#C9B458")

    default_font = tkFont.Font(family="Arial", size=14, weight="bold")
    root.option_add("*Font", default_font)
    root.option_add('*justify', 'right')

    icon_path = os.path.join("images", "ITP1.ICO")
    if os.path.exists(icon_path):
        root.iconbitmap(icon_path)

    def on_close():
        messagebox.showinfo("تنبيه ⚠️", "استعمل زر العودة للقائمة الرئيسية")
    root.protocol("WM_DELETE_WINDOW", on_close)

    total1 = total2 = total3 = 0.0
    all_data = {}

    if os.path.exists("data1.json"):
        with open("data1.json", "r", encoding="utf-8-sig") as f:
            try:
                all_data = json.load(f)
                total1 = float(all_data.get("PV2_2", {}).get("total", 0))
                total2 = float(all_data.get("PV2_3", {}).get("total1", 0))
                total3 = float(all_data.get("PV2_5", {}).get("total2", 0))
            except Exception as e:
                print("خطأ تحميل JSON:", e)

    td = round(total1 + total2 + total3, 2)
    moyenne = round(td / 3, 2)
    decision = "ناجح" if moyenne >= 10 else "راسب"
    today = datetime.now().strftime("%Y-%m-%d")

    # ---------- الإطار العلوي ----------
    top_frame = tk.Frame(root, bg="#C9B458")
    top_frame.pack(pady=15)

    labels = ["الإختبار الثالث", "الإختبار الثاني", "الإختبار الأول", "التاريخ"]
    values = [total3, total2, total1, today]

    for i in range(4):
        tk.Label(top_frame, text=labels[i], bg="#C9B458", fg="#2F2F2F",
                 width=20, relief="ridge").grid(row=0, column=i, padx=6)

        e = tk.Entry(top_frame, width=20, justify="center",
                     bg="white", fg="#2F2F2F", relief="solid", bd=1)
        e.grid(row=1, column=i, padx=6)
        e.insert(0, f"{values[i]:.2f}" if isinstance(values[i], (int, float)) else values[i])
        e.config(state="readonly")

    row2_labels = ["القرار", "المعدل", "المجموع"]
    row2_values = [decision, moyenne, td]

    for i in range(3):
        tk.Label(top_frame, text=row2_labels[i], bg="#C9B458", fg="#2F2F2F",
                 width=20, relief="ridge").grid(row=2, column=i, padx=6, pady=(10, 0))

        e = tk.Entry(top_frame, width=20, justify="center",
                     bg="white", fg="#2F2F2F", relief="solid", bd=1)
        e.grid(row=3, column=i, padx=6)
        e.insert(0, row2_values[i])
        e.config(state="readonly")

    # ---------- جدول الأعضاء ----------
    table_frame = tk.Frame(root, bg="#C9B458")
    table_frame.pack(pady=25)

    headers = ["الوظيفة", "الاسم", "اللقب", "العضو"]
    for col, h in enumerate(headers):
        tk.Label(table_frame, text=h, bg="#C9B458", fg="#2F2F2F",
                 width=22, relief="ridge").grid(row=0, column=col, padx=3)

    members = ["مفتش المقاطعة", "العضو الأول", "العضو الثاني"]
    entries_job, entries_fname, entries_lname = [], [], []

    for i in range(len(members)):
        job = tk.Entry(table_frame, width=25, justify="center", bg="white", relief="solid", bd=1)
        fname = tk.Entry(table_frame, width=25, justify="center", bg="white", relief="solid", bd=1)
        lname = tk.Entry(table_frame, width=25, justify="center", bg="white", relief="solid", bd=1)

        job.grid(row=i + 1, column=0, padx=3, pady=6)
        fname.grid(row=i + 1, column=1, padx=3, pady=6)
        lname.grid(row=i + 1, column=2, padx=3, pady=6)

        tk.Label(table_frame, text=members[i], bg="#EFEBD8",
                 width=15, relief="ridge").grid(row=i + 1, column=3, padx=3, pady=6, ipady=6)

        entries_job.append(job)
        entries_fname.append(fname)
        entries_lname.append(lname)

    # ---------- الأزرار ----------
    buttons_frame = tk.Frame(root, bg="#C9B458")
    buttons_frame.pack(fill="x", padx=20, pady=25)

    def save_data():
        all_data["PV2_4"] = {
            "total": td,
            "average": moyenne,
            "decision": decision,
            "date": today,
            "members": []
        }

        for i in range(len(members)):
            all_data["PV2_4"]["members"].append({
                "member": members[i],
                "job": entries_job[i].get(),
                "fname": entries_fname[i].get(),
                "lname": entries_lname[i].get()
            })

        save_to_json(all_data)
        messagebox.showinfo("تم الحفظ ✅", "تم حفظ البيانات بنجاح")
        fill_button.config(state="normal")

    def fill_all_templates():
        notes_win = tk.Toplevel(root)
        notes_win.title("الملاحظات والتقدير العام")
        notes_win.geometry("600x350")
        notes_win.configure(bg="#C9B458")

        if os.path.exists(icon_path):
            notes_win.iconbitmap(icon_path)

        tk.Label(notes_win, text="الملاحظات والتقديرات العامة",
                 bg="#C9B458", fg="#2F2F2F",
                 width=50, relief="ridge").pack(pady=10)

        notes_text = tk.Text(notes_win, height=8, width=65, wrap="word", bg="white")
        notes_text.pack(pady=10)

        def save_notes_and_fill():
            notes = notes_text.get("1.0", tk.END).strip()
            if not notes:
                messagebox.showerror("خطأ", "يرجى إدخال الملاحظات")
                return

            all_data.setdefault("PV2_4", {})["notes"] = notes
            save_to_json(all_data)

            template_path = "templates/ADJOIN.docx"
            if not os.path.exists(template_path):
                messagebox.showerror("خطأ", "القالب غير موجود")
                return

            adjoin_last_name = all_data.get("ADJOIN", {}).get("last_name", "بدون_اسم")
            doc = Document(template_path)

            # 🔴 التغيير الأساسي: دمج كل الأقسام في جذر القالب
            flat_data = flatten_all_sections(all_data)
            replace_in_doc(doc, flat_data)

            output_dir = r"C:\Users\TechSpace\Documents\محاضر التثبيت\مساعد التكوين"
            os.makedirs(output_dir, exist_ok=True)

            output_path = os.path.join(output_dir, f"PV2_{adjoin_last_name}.docx")
            doc.save(output_path)

            try:
                convert(output_path, output_path.replace(".docx", ".pdf"))
            except:
                pass

            messagebox.showinfo("تم بنجاح ✅", "تم حفظ البيانات وتعبئة المحضر")
            notes_win.destroy()

        tk.Button(
            notes_win,
            text="💾 حفظ الملاحظات وتعبئة المحضر",
            bg="#1E88E5", fg="white",
            width=30,
            command=save_notes_and_fill
        ).pack(pady=15)

    tk.Button(
        buttons_frame,
        text="💾 حفظ البيانات",
        bg="#4CAF50", fg="white",
        width=22,
        command=save_data
    ).pack(side="left", expand=True)

    fill_button = tk.Button(
        buttons_frame,
        text="🧾 تعبئة المحضر",
        bg="#1E88E5", fg="white",
        width=22,
        state="disabled",
        command=fill_all_templates
    )
    fill_button.pack(side="left", expand=True)

    tk.Button(
        buttons_frame,
        text="⬅️ العودة للقائمة الرئيسية",
        bg="#C0392B", fg="white",
        width=22,
        command=lambda: (root.destroy(), ADJOIN.open_confirmation_window())
    ).pack(side="left", expand=True)

    root.mainloop()

if __name__ == "__main__":
    open_pv2_4()
