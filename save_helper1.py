import json
import os
from datetime import datetime
from docx import Document
import re

# ======================================================
# ✅ دالة حفظ البيانات في ملف JSON (آمنة + UTF-8-SIG)
# ======================================================
def save_to_json(data, file_path="data.json"):
    """
    تحفظ البيانات في ملف JSON باستخدام ترميز utf-8-sig
    للحفاظ على الحروف العربية بشكل سليم في Windows وWord.
    """
    try:
        try:
            with open(file_path, "r", encoding="utf-8-sig") as file:
                all_data = json.load(file)
        except FileNotFoundError:
            all_data = {}

        all_data.update(data)

        with open(file_path, "w", encoding="utf-8-sig") as file:
            json.dump(all_data, file, ensure_ascii=False, indent=4)

        print(f"✅ تم حفظ البيانات في {file_path}")
    except Exception as e:
        print(f"❌ خطأ أثناء حفظ البيانات في JSON: {e}")


# ======================================================
# ✅ دالة ملء قالب Word بالبيانات بشكل ذكي
# ======================================================
def fill_template_with_check(data):
    """
    تملأ قالب Word بالبيانات من JSON أو قاموس Python.
    - تدعم القوالب التي تحتوي على placeholders من نوع {{placeholder}}
    - تتأكد من وجود كل المفاتيح
    - تحفظ الملف في مجلد خاص على سطح المكتب
    """
    try:
        # --- 1️⃣ قراءة البيانات ---
        if isinstance(data, dict):
            json_data = data
        else:
            with open(data, "r", encoding="utf-8-sig") as f:
                json_data = json.load(f)

        teacher = json_data.get("Teacher", {})
        inspector = json_data.get("Inspector", {})

        replacements = {}
        replacements.update(teacher)
        replacements.update(inspector)

        # --- 2️⃣ تحديد القالب ---
        template_path = r"C:\Users\TechSpace\Desktop\Nouveau projet4\Nouveau dossier\templates\PFP.docx"
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"❌ لم يتم العثور على القالب في {template_path}")

        doc = Document(template_path)
        placeholders_in_doc = set()

        # --- 3️⃣ استخراج جميع الـ placeholders ---
        def extract_placeholders(text):
            matches = re.findall(r"\{\{(.*?)\}\}", text)
            for m in matches:
                placeholders_in_doc.add(m.strip())

        for p in doc.paragraphs:
            extract_placeholders(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        extract_placeholders(p.text)

        print(f"📄 جميع placeholders في القالب: {placeholders_in_doc}")

        # --- 4️⃣ التأكد من وجود جميع المفاتيح ---
        for key in placeholders_in_doc:
            if key not in replacements:
                replacements[key] = ""

        # --- 5️⃣ استبدال النصوص ---
        def replace_placeholder_in_paragraph(paragraph):
            full_text = "".join(run.text for run in paragraph.runs)
            for key, value in replacements.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(value))
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = full_text
            else:
                paragraph.add_run(full_text)

        for p in doc.paragraphs:
            replace_placeholder_in_paragraph(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_placeholder_in_paragraph(p)

        # --- 6️⃣ حفظ الملف الناتج على سطح المكتب ---
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        year = datetime.now().year
        output_folder = os.path.join(desktop_path, f"محاضر التثبيت لسنة {year}")
        os.makedirs(output_folder, exist_ok=True)

        teacher_name = f"{teacher.get('last_name','')}_{teacher.get('first_name','')}".strip("_")
        if not teacher_name:
            teacher_name = "أستاذ_غير_معروف"

        output_path = os.path.join(output_folder, f"محضر التثبيت {teacher_name}_{year}.docx")
        if os.path.exists(output_path):
            os.remove(output_path)

        doc.save(output_path)
        print(f"✅ تم إنشاء الملف بنجاح في: {output_path}")
        return output_path

    except Exception as e:
        print(f"❌ خطأ أثناء ملء القالب: {e}")
        return None
# ======================================================
# ✅ دالة بديلة لضمان التوافق مع الملفات القديمة
# ======================================================
def save_to_json_shared(form_name, data, file_path=None):
    """
    دالة توافقية: تحفظ بيانات الفورم تحت مفتاح اسم الفورم داخل ملف JSON مشترك.
    - file_path: مسار الملف إن أردت (افتراضيًا form_data.json داخل نفس المجلد).
    """
    try:
        if file_path is None:
            file_path = os.path.join(os.path.dirname(__file__), "form_data.json")

        # تأكد أن البيانات dict
        if not isinstance(data, dict):
            print("⚠️ البيانات المرسلة إلى save_to_json_shared ليست من نوع dict.")
            return

        # حمّل الموجود أو أنشئ هيكل جديد
        try:
            if os.path.exists(file_path):
                with open(file_path, "r", encoding="utf-8-sig") as f:
                    all_data = json.load(f)
            else:
                all_data = {}
        except Exception:
            all_data = {}

        # حدّث قسم الفورم
        all_data[form_name] = data

        # اكتب الملف
        with open(file_path, "w", encoding="utf-8-sig") as f:
            json.dump(all_data, f, ensure_ascii=False, indent=4)

        print(f"✅ تم حفظ البيانات المشتركة للفورم: {form_name} في {file_path}")
    except Exception as e:
        print(f"❌ خطأ أثناء الحفظ باستخدام save_to_json_shared: {e}")


# ======================================================
# ✅ دالة تحميل بيانات فورم محفوظة (مساعدة)
# ======================================================
def load_from_json_shared(form_name, file_path=None):
    """
    إرجاع البيانات المحفوظة تحت مفتاح form_name إن وُجدت، وإلا None.
    """
    try:
        if file_path is None:
            file_path = os.path.join(os.path.dirname(__file__), "form_data.json")
        if not os.path.exists(file_path):
            return None
        with open(file_path, "r", encoding="utf-8-sig") as f:
            all_data = json.load(f)
        return all_data.get(form_name)
    except Exception as e:
        print(f"❌ خطأ أثناء load_from_json_shared: {e}")
        return None
